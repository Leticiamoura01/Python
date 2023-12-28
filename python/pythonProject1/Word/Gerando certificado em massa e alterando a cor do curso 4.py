from docx import Document
from docx.shared import Pt

from docx.shared import RGBColor

from openpyxl import load_workbook

#Pegando o caminho do arquivo
nome_arquivo_alunos = "DadosAluno.xlsx"
planilhaDadosAlunos = load_workbook(nome_arquivo_alunos)

#Selecionando a planilha/sheet/aba
sheet_selecionada = planilhaDadosAlunos["Nomes"]

for linha in range(2, len(sheet_selecionada["A"]) + 1):

        #Abre o arquivo do word
        arquivoWord = Document("C:\\Users\kauan\Downloads\\1 rpa\python\\pythonProject1\\Word\\Certificado3.docx")

        #Seleciona o estilo
        estilo = arquivoWord.styles["Normal"]

        #Pegando o nome do aluno quando passar na celula
        nomeAluno = sheet_selecionada['A%s' % linha].value
        dia = sheet_selecionada['B%s' % linha].value
        mes = sheet_selecionada['C%s' % linha].value
        ano = sheet_selecionada['D%s' % linha].value
        nomeCurso = sheet_selecionada['E%s' % linha].value
        nomeInstrutor = sheet_selecionada['F%s' % linha].value

        #for = para
        for paragrafo in arquivoWord.paragraphs:

            #if = se
            if "@nome" in paragrafo.text:
                paragrafo.text = nomeAluno
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            paragrafoP1 = "Concluiu com sucesso o curso de"
            paragrafoP2 = ", como carga horaria de 20 horas"
            paragrafoComplemento = f"{paragrafoP2} {dia} de {mes} de {ano}."

            if "escola" in paragrafo.text:
                paragrafo.text = paragrafoP1
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)
                adicionaNovaPalavra = paragrafo.add_run(nomeCurso)
                adicionaNovaPalavra.font.color.rgb = RGBColor(255, 0, 0) #Texto vermelho
                adicionaNovaPalavra.underline = True #Sublinhado
                adicionaNovaPalavra.bold = True #Negrito
                adicionaNovaPalavra = paragrafo.add_run(paragrafoComplemento)
                adicionaNovaPalavra.font.color.rgb = RGBColor(0, 0, 0)


            if "instrutor" in paragrafo.text:
                paragrafo.text = nomeInstrutor + "- Instrutor"
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        caminhoCertificado = "coloca o caminho dapasta criada so pra certificado"

        #Salvando o certificado do aluno
        arquivoWord.save(caminhoCertificado)

print("Certificados gerados com sucesso!")

